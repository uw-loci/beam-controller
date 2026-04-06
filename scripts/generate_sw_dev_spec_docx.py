from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "BCON SW Dev Spec _2026-03-28B-updated.docx"


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_document() -> Document:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.add_run("BCON Software Development Specification\n")
    title.add_run("Rev 2026-03-28B")

    intro = document.add_paragraph()
    intro.add_run("Purpose: ").bold = True
    intro.add_run(
        "Define the current software behavior, interfaces, and safety constraints for the "
        "BCON Arduino Mega pulser-control firmware and its host-side integration."
    )

    scope = document.add_paragraph()
    scope.add_run("Scope note: ").bold = True
    scope.add_run(
        "This revision replaces older dashboard/deflection-centric descriptions with the "
        "current implementation actually present in BCON_mega_modbus.ino and the repository README."
    )

    document.add_heading("1. Intended Purpose and Scope", level=1)
    add_bullets(
        document,
        [
            "The BCON microcontroller firmware turns an Arduino Mega 2560 into a 3-channel beam-gate pulser controller.",
            "The firmware owns pulser gate outputs, pulser enable-toggle outputs, Modbus RTU communications, watchdog enforcement, interlock handling, and local LCD diagnostics.",
            "Beam deflection waveform generation, Kepco BOP control, LUT interpolation, beam position tracking, and physics calculations are out of scope for this firmware and belong to other subsystems.",
            "The firmware is safety-oriented: any non-Ready top-level state forces all gate outputs LOW.",
        ],
    )

    document.add_heading("2. Software Boundary", level=1)
    add_table(
        document,
        ["Area", "Owned by This Firmware", "Notes"],
        [
            ["Pulser output control", "Yes", "Per-channel Off/DC/Pulse/PulseTrain control"],
            ["Watchdog safety", "Yes", "Software watchdog plus AVR hardware WDT"],
            ["Knob Box interlock readback", "Yes", "Active-HIGH interlock input on A0"],
            ["Fault latch status/reporting", "Partial", "Fault latch register path exists; hardware sensing path is not fully implemented in the current .ino"],
            ["Kepco waveform generation", "No", "Out of scope for the Arduino firmware"],
            ["Beam LUT calculations", "No", "Handled outside the microcontroller"],
            ["Dashboard GUI", "No", "Host-side Python subsystem"],
        ],
    )

    document.add_heading("3. Hardware and Runtime Platform", level=1)
    add_bullets(
        document,
        [
            "Target MCU: Arduino Mega 2560 (ATmega2560).",
            "Firmware language: Arduino C++.",
            "Primary communication stack: Modbus RTU using the modbus-esp8266 ModbusRTU library.",
            "Optional local diagnostics: 20x4 I2C LCD via LiquidCrystal_I2C.",
            "Runtime recovery: AVR hardware watchdog enabled with an 8 second timeout and disabled early during .init3 after reset.",
        ],
    )

    document.add_heading("4. Build-Time Configuration", level=1)
    add_table(
        document,
        ["Macro", "Current Value", "Meaning"],
        [
            ["BCON_USE_USB_SERIAL", "1", "Bench/debug mode via USB CDC Serial. Set to 0 for production RS-485 on Serial1."],
            ["BCON_ENABLE_LCD", "1", "Enable 20x4 I2C LCD diagnostics."],
        ],
    )

    document.add_heading("5. Pin Assignments", level=1)
    add_table(
        document,
        ["Signal", "Pin", "Direction", "Description"],
        [
            ["INTERLOCK", "A0", "Input", "Active-HIGH knob-box interlock feedback"],
            ["GATE_A", "A11", "Output", "Channel A gate drive"],
            ["GATE_B", "A13", "Output", "Channel B gate drive"],
            ["GATE_C", "A9", "Output", "Channel C gate drive"],
            ["LED_A", "D3", "Output", "Channel A gate status LED"],
            ["LED_B", "D5", "Output", "Channel B gate status LED"],
            ["LED_C", "D7", "Output", "Channel C gate status LED"],
            ["ENA_A", "D12", "Output", "Channel A pulser enable-toggle"],
            ["ENA_B", "D11", "Output", "Channel B pulser enable-toggle"],
            ["ENA_C", "D10", "Output", "Channel C pulser enable-toggle"],
            ["RS485_DE_RE", "D17", "Output", "Direction control when RS-485 mode is enabled"],
            ["I2C SDA/SCL", "D20/D21", "Bidirectional", "LCD bus via Wire"],
        ],
    )

    document.add_heading("6. Top-Level Safety State Machine", level=1)
    add_table(
        document,
        ["State", "Code", "Meaning"],
        [
            ["Ready", "0", "Outputs may operate normally"],
            ["SafeInterlock", "1", "Interlock not satisfied; outputs forced LOW"],
            ["SafeWatchdog", "2", "Host heartbeat timeout; outputs forced LOW"],
            ["FaultLatched", "3", "Fault latched; outputs forced LOW"],
        ],
    )
    add_bullets(
        document,
        [
            "State priority is evaluated in this order: Interlock, Watchdog, FaultLatched, Ready.",
            "Any state other than Ready suppresses all gate outputs.",
            "A mode write for DC/Pulse/PulseTrain is rejected unless evaluateState() returns Ready at callback time.",
        ],
    )

    document.add_heading("7. Watchdog Design", level=1)
    add_bullets(
        document,
        [
            "Software watchdog register: WATCHDOG_MS at holding register 0.",
            "Default timeout: 1500 ms. Allowed range: 50 to 60000 ms.",
            "Boot grace period: 8000 ms before software watchdog enforcement starts.",
            "Writes feed the watchdog through all ON_SET callbacks.",
            "Reads of SYS_STATE (register 100) also feed the watchdog, matching the host polling model.",
            "Hardware watchdog: AVR WDT enabled at the end of setup() with WDTO_8S and reset in every loop() iteration.",
            "Early startup hook in .init3 captures MCUSR and disables WDT to avoid reset-loop behavior after a watchdog reset.",
        ],
    )

    document.add_heading("8. Output Modes", level=1)
    add_table(
        document,
        ["Mode", "Code", "Behavior"],
        [
            ["Off", "0", "Gate LOW, channel idle"],
            ["DC", "1", "Gate held HIGH continuously while system is Ready"],
            ["Pulse", "2", "Single pulse if COUNT equals 1"],
            ["PulseTrain", "3", "Repeated HIGH/LOW pulse train"],
        ],
    )
    add_bullets(
        document,
        [
            "If mode 2 (Pulse) is written while COUNT > 1, the firmware elevates the channel to PulseTrain automatically.",
            "Pulse duration range is 1 to 60000 ms.",
            "Pulse count range is 1 to 10000.",
            "Pulse-train gap equals pulse duration, resulting in a 50 percent duty cycle waveform.",
        ],
    )

    document.add_heading("9. Modbus RTU Interface", level=1)
    add_bullets(
        document,
        [
            "Slave ID: 1.",
            "Baud: 115200, 8N1.",
            "Bench mode uses Serial when BCON_USE_USB_SERIAL = 1.",
            "Production mode uses Serial1 with DE/RE on D17 when BCON_USE_USB_SERIAL = 0.",
            "Supported function codes are 0x03, 0x06, and 0x10 via the ModbusRTU library.",
        ],
    )

    document.add_heading("10. Holding Register Map", level=1)
    add_table(
        document,
        ["Address", "Name", "R/W", "Meaning"],
        [
            ["0", "WATCHDOG_MS", "R/W", "Software watchdog timeout in milliseconds"],
            ["1", "TELEMETRY_MS", "R/W", "Informational poll period value"],
            ["2", "COMMAND", "W", "NOP, all-off, clear-fault commands"],
            ["10/20/30", "CHx_MODE", "R/W", "Per-channel output mode"],
            ["11/21/31", "CHx_PULSE_MS", "R/W", "Per-channel pulse duration"],
            ["12/22/32", "CHx_COUNT", "R/W", "Per-channel pulse count"],
            ["13/23/33", "CHx_ENABLE_TOGGLE", "W", "Write 1 for 100 ms enable-toggle pulse"],
            ["100", "SYS_STATE", "R", "Current top-level state and watchdog keepalive read path"],
            ["101", "SYS_REASON", "R", "Reason/status mirror"],
            ["102", "FAULT_LATCHED", "R", "Fault latch status"],
            ["103", "INTERLOCK_OK", "R", "Interlock status"],
            ["104", "WATCHDOG_OK", "R", "Watchdog status"],
            ["105", "LAST_ERROR", "R", "Last error register"],
            ["110-118", "CH1 status", "R", "Mode, timing, count, output level"],
            ["120-128", "CH2 status", "R", "Mode, timing, count, output level"],
            ["130-138", "CH3 status", "R", "Mode, timing, count, output level"],
        ],
    )

    document.add_heading("11. Command Register Semantics", level=1)
    add_table(
        document,
        ["COMMAND Value", "Action"],
        [
            ["0", "NOP; may be used as a heartbeat write"],
            ["1", "All channels off"],
            ["2 or 3", "Clear fault latch when allowed by system state"],
        ],
    )

    document.add_heading("12. Host Integration Expectations", level=1)
    add_bullets(
        document,
        [
            "A host should wait for the Arduino reset/re-enumeration interval before beginning Modbus traffic in USB bench mode.",
            "The host should poll system status registers 100 through 105 and channel status blocks 110-118, 120-128, and 130-138.",
            "Either periodic writes or SYS_STATE reads keep the software watchdog alive.",
            "The current Python GUI additionally sends a periodic COMMAND=0 write as a defense-in-depth heartbeat.",
        ],
    )

    document.add_heading("13. Current Implementation Limitations", level=1)
    add_bullets(
        document,
        [
            "The current firmware exposes FAULT_LATCHED and related semantics, but the overcurrent sensing path itself is not present in the current source and must be implemented if required by hardware.",
            "Per-channel external status inputs described in older documents are not wired into the current firmware; status offsets +4 through +7 currently return zero.",
            "README text in the repo still contains one stale statement claiming modbus-esp8266 is unused; the firmware does use ModbusRTU from that library.",
        ],
    )

    document.add_heading("14. Differences From Older SW Dev Spec Revisions", level=1)
    add_bullets(
        document,
        [
            "Older revisions described dashboard waveform generation, Kepco BOP control, LUT files, and beam-deflection calculations as if they were part of the Arduino firmware. They are not part of the current microcontroller implementation.",
            "Older revisions used an obsolete register map with COMMAND at register 0 and DAC/waveform registers immediately after. That register map is not valid for the current firmware.",
            "This revision should be treated as the authoritative MCU-side software description unless and until the firmware changes again.",
        ],
    )

    document.add_heading("15. Recommended Verification", level=1)
    add_bullets(
        document,
        [
            "Verify read access to 0, 1, 2, 100-105, and 110-138 against the live device.",
            "Verify pulse, pulse-train, and all-off behavior on all three channels.",
            "Verify watchdog hold-in using both SYS_STATE polling and COMMAND=0 writes.",
            "Verify a rejected mode write while interlock is low or watchdog is expired returns NotReady behavior without recursion or firmware reset.",
        ],
    )

    document.add_section(WD_SECTION_START.NEW_PAGE)
    appendix = document.add_paragraph()
    appendix.add_run("Generated from repository state on 2026-03-28. ").bold = True
    appendix.add_run(
        "Primary implementation references: BCON_mega_modbus/BCON_mega_modbus.ino, README.md, and the reviewed prior SW Dev Spec."
    )

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()